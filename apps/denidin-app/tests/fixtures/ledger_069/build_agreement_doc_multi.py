"""Build `agreement_doc_multi.docx` (Feature 069 US10 fixture).

Deterministic generator so the binary can be regenerated from source and
reviewed as text. Run from `apps/denidin-app` with its venv active:

    python3 tests/fixtures/ledger_069/build_agreement_doc_multi.py

Keep the paragraph/table text in exact sync with
`agreement_doc_multi.manifest.json`.
"""
from pathlib import Path

from docx import Document

OUT = Path(__file__).with_name("agreement_doc_multi.docx")

PARAS = [
    "הסכם שכר טרחה",
    "",
    "נחתם ביום 22/05/2026 בין משרד עורכי הדין (\"המשרד\") לבין הלקוחה גב' נטלי ברנשטיין "
    "(ת.ז. 041228876), מרחוב האלון 7, רעננה (\"הלקוחה\").",
    "",
    "המשרד ייצג את הלקוחה בתביעה הכספית. שכר הטרחה המוסכם:",
]

TABLE_ROWS = [
    ("רכיב", "תיאור", "סכום / שיעור"),
    ("מקדמה קבועה", "לתשלום עם חתימת ההסכם", '18,000 ש"ח + מע"מ'),
    ("שכר הצלחה", "מכל סכום שייפסק או ייגבה לטובת הלקוחה", '18% + מע"מ'),
    ("שכר עבור דיון", "עבור כל ישיבת הוכחות בבית המשפט", '2,000 ש"ח + מע"מ'),
]

TAIL = [
    "",
    "התשלומים ישולמו על ידי בעלה של הלקוחה, מר דניאל ברנשטיין, מבלי לגרוע מאחריות הלקוחה.",
    "",
    "חתימת הלקוחה: ____________     חתימת המשרד: ____________",
]


def main() -> None:
    doc = Document()
    for p in PARAS:
        doc.add_paragraph(p)
    table = doc.add_table(rows=0, cols=3)
    for row in TABLE_ROWS:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
    for p in TAIL:
        doc.add_paragraph(p)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
