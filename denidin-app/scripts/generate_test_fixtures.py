"""
Generate test fixture documents (PDF, DOCX) from text templates.

Run this script to create test documents:
    python scripts/generate_test_fixtures.py

For PDFs with Hebrew: Uses PIL to create image-based PDFs with proper Hebrew rendering.
"""

from pathlib import Path
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import io


def create_docx_from_text(text_content: str, output_path: Path):
    """Create a DOCX file from text content."""
    doc = Document()
    
    # Add title
    doc.add_heading('מסמך', 0)  # Hebrew title
    
    # Add content (preserve line breaks)
    for line in text_content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            # Right-to-left for Hebrew
            p.paragraph_format.left_to_right = False
        else:
            doc.add_paragraph()  # Empty paragraph for spacing
    
    doc.save(output_path)
    print(f"✅ Created: {output_path}")


def create_pdf_from_text_as_image(text_content: str, output_path: Path):
    """
    Create a PDF from text by rendering it as an image.
    This properly handles Hebrew text rendering.
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    
    # Create image with Hebrew text
    width, height = 800, 1100  # A4-ish proportions
    img = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use Arial font (has Hebrew support)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 28)
    except:
        # Fallback to default (won't render Hebrew properly)
        font = ImageFont.load_default()
        print("⚠️  Warning: Arial font not found, Hebrew may not display correctly")
    
    # Draw text line by line (RTL - from right)
    y = 50
    line_height = 40
    
    for line in text_content.split('\n'):
        if line.strip():
            # Draw from right for Hebrew (RTL)
            draw.text((width - 50, y), line, fill='black', font=font, anchor='rt')
        y += line_height
    
    # Save image temporarily
    temp_img_path = output_path.parent / f"{output_path.stem}_temp.png"
    img.save(temp_img_path, format='PNG')
    
    # Create PDF with the image
    c = canvas.Canvas(str(output_path), pagesize=A4)
    pdf_width, pdf_height = A4
    
    # Draw saved image on PDF page
    c.drawImage(str(temp_img_path), 0, 0, width=pdf_width, height=pdf_height)
    c.save()
    
    # Clean up temp file
    temp_img_path.unlink()
    
    print(f"✅ Created: {output_path}")


def main():
    """Generate all test fixture documents."""
    fixtures_dir = Path(__file__).parent.parent / "tests" / "fixtures" / "media"
    
    print("Generating test fixture documents...")
    print(f"Directory: {fixtures_dir}\n")
    
    # 1. Contract (PDF + DOCX)
    contract_txt = fixtures_dir / "contract_peter_adam.txt"
    if contract_txt.exists():
        content = contract_txt.read_text(encoding='utf-8')
        
        # Create PDF (as image for Hebrew support)
        create_pdf_from_text_as_image(content, fixtures_dir / "contract_peter_adam.pdf")
        
        # Create DOCX
        create_docx_from_text(content, fixtures_dir / "contract_peter_adam.docx")
    else:
        print(f"❌ Missing: {contract_txt}")
    
    # 2. Generic document (PDF)
    generic_txt = fixtures_dir / "document_no_client.txt"
    if generic_txt.exists():
        content = generic_txt.read_text(encoding='utf-8')
        
        # Create PDF (as image for Hebrew support)
        create_pdf_from_text_as_image(content, fixtures_dir / "document_no_client.pdf")
    else:
        print(f"❌ Missing: {generic_txt}")
    
    print("\n✅ Test fixture generation complete!")
    print("\nManual steps required:")
    print("1. Add a real receipt photo as: receipt_cafe.jpg")
    print("2. Add any audio file as: unsupported.mp3")
    print("\nYou can also manually create better PDFs by:")
    print("- Opening .txt files in Word/Pages")
    print("- Saving as PDF (better Hebrew rendering)")


if __name__ == "__main__":
    main()
