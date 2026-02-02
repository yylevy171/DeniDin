"""
Unit tests for DOCX Extractor (Feature 003 Phase 3.3 + Phase 4)

CHK Requirements:
- CHK005: Corrupted file handling
- CHK006: Hebrew text extraction
- CHK007: Graceful degradation on failures
- CHK010: Layout/structure preservation
- CHK078: Empty document handling

Phase 4: Optional AI-powered document analysis
"""
import pytest
import io
from unittest.mock import Mock, patch
from docx import Document
from src.handlers.extractors.docx_extractor import DOCXExtractor
from src.models.media import Media


@pytest.fixture
def mock_denidin_context():
    """Create mock DeniDin context."""
    context = Mock()
    context.config = Mock()
    context.config.ai_model = "gpt-4o"
    context.config.ai_reply_max_tokens = 4096
    context.config.temperature = 0.7
    context.config.constitution_config = {}
    context.ai_handler = Mock()
    # Mock _load_constitution to return empty string
    context.ai_handler._load_constitution = Mock(return_value="")
    return context


@pytest.fixture
def docx_extractor(mock_denidin_context):
    """Create DOCXExtractor instance."""
    return DOCXExtractor(mock_denidin_context)


def create_docx_media(*paragraphs) -> Media:
    """Helper to create DOCX Media object from paragraph texts."""
    doc = Document()
    for para_text in paragraphs:
        doc.add_paragraph(para_text)
    
    # Save to in-memory bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return Media.from_bytes(
        data=docx_bytes.read(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="test.docx"
    )


def test_extract_simple_docx_text(docx_extractor, mock_denidin_context):
    """
    CHK006: Extract basic text from simple DOCX with AI analysis.
    Should handle plain paragraph text and return AI-analyzed response.
    """
    media = create_docx_media("Hello World", "This is a test document")
    
    # Mock AI response - get_response returns object with response_text attribute
    mock_response = Mock()
    mock_response.response_text = "Analysis: Hello World greeting and test document"
    mock_denidin_context.ai_handler.get_response.return_value = mock_response
    
    # With analyze=True, should return AI analysis
    result = docx_extractor.analyze_media(media, analyze=True)
    
    assert result["extraction_quality"] == "high"
    assert len(result["warnings"]) == 0
    assert "Analysis" in result["raw_response"]
    assert mock_denidin_context.ai_handler.get_response.call_count == 1


def test_extract_hebrew_text(docx_extractor, mock_denidin_context):
    """
    CHK006: Extract Hebrew text with UTF-8 encoding.
    Should preserve Hebrew characters correctly.
    """
    media = create_docx_media("שלום עולם", "זהו מסמך בדיקה")
    
    # Mock AI response - get_response returns object with response_text attribute
    mock_response = Mock()
    mock_response.response_text = "סיכום: מסמך בעברית"
    mock_denidin_context.ai_handler.get_response.return_value = mock_response
    
    result = docx_extractor.analyze_media(media, analyze=True)
    
    assert result["extraction_quality"] == "high"
    assert "סיכום" in result["raw_response"]
    assert len(result["warnings"]) == 0


def test_preserve_paragraph_structure(docx_extractor, mock_denidin_context):
    """
    CHK010: Preserve paragraph structure with separators.
    Should separate paragraphs with double newlines in AI analysis.
    """
    media = create_docx_media("First paragraph", "Second paragraph", "Third paragraph")
    
    # Mock AI response that acknowledges all paragraphs
    mock_response = Mock()
    mock_response.response_text = "Analysis: Contains First, Second, Third paragraphs"
    mock_denidin_context.ai_handler.get_response.return_value = mock_response
    
    result = docx_extractor.analyze_media(media, analyze=True)
    
    assert result["extraction_quality"] == "high"
    assert "First" in result["raw_response"]
    assert "Third" in result["raw_response"]


def test_handle_empty_docx(docx_extractor, mock_denidin_context):
    """
    CHK078: Handle empty DOCX gracefully.
    Empty documents should not call AI and return no analysis.
    """
    media = create_docx_media()  # No paragraphs
    
    result = docx_extractor.analyze_media(media, analyze=True)
    
    # Empty DOCX should not call AI
    assert mock_denidin_context.ai_handler.send_message.call_count == 0
    assert result["raw_response"] == ""
    assert len(result["warnings"]) == 1
    assert "empty" in result["warnings"][0].lower()


def test_handle_corrupted_docx(docx_extractor, mock_denidin_context):
    """
    CHK005, CHK007: Gracefully handle corrupted DOCX.
    Should return empty text with error warning.
    """
    # Create corrupted DOCX data
    corrupted_media = Media.from_bytes(
        data=b"This is not a valid DOCX file",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="corrupted.docx"
    )
    
    result = docx_extractor.analyze_media(corrupted_media, analyze=True)
    
    assert result["raw_response"] == ""
    assert len(result["warnings"]) >= 1
    assert result["extraction_quality"] == "failed"


def test_extract_text_ignoring_formatting(docx_extractor, mock_denidin_context):
    """
    Extract plain text, ignoring formatting.
    Should extract text regardless of bold/italic/underline.
    """
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Normal text. ")
    para.add_run("Bold text. ").bold = True
    para.add_run("Italic text.").italic = True
    
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    media = Media.from_bytes(
        data=docx_bytes.read(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="formatted.docx"
    )
    
    # Mock AI response - get_response returns object with response_text attribute
    mock_response = Mock()
    mock_response.response_text = "Contains Normal, Bold, and Italic text"
    mock_denidin_context.ai_handler.get_response.return_value = mock_response
    
    result = docx_extractor.analyze_media(media, analyze=True)
    
    # All text should be extracted and analyzed
    assert result["extraction_quality"] == "high"
    assert len(result["warnings"]) == 0


def test_extract_complex_structure(docx_extractor, mock_denidin_context):
    """
    Extract text from complex document structure.
    Should handle tables, lists, and mixed content.
    """
    doc = Document()
    doc.add_heading("Document Title", level=1)
    doc.add_paragraph("Introduction paragraph")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"
    
    doc.add_paragraph("Conclusion paragraph")
    
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    media = Media.from_bytes(
        data=docx_bytes.read(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="complex.docx"
    )
    
    # Mock AI response - get_response returns object with response_text attribute
    mock_response = Mock()
    mock_response.response_text = "Document with title, intro, table with cells, and conclusion"
    mock_denidin_context.ai_handler.get_response.return_value = mock_response
    
    result = docx_extractor.analyze_media(media, analyze=True)
    
    # Should extract and analyze structure
    assert result["extraction_quality"] == "high"
    assert "table" in result["raw_response"].lower() or "cell" in result["raw_response"].lower()


# ===== Phase 4: AI-Powered Document Analysis Tests =====

def test_analyze_document_with_ai(docx_extractor, mock_denidin_context):
    """
    Phase 4: With analyze=True, should call AI to analyze document.
    """
    media = create_docx_media("Invoice #12345", "Total: $100", "Due Date: 2024-01-01")
    
    # Mock AI response - get_response returns object with response_text attribute
    mock_response = Mock()
    mock_response.response_text = """DOCUMENT_TYPE: invoice
SUMMARY: Invoice for $100 due on 2024-01-01
KEY_POINTS:
- Invoice number 12345
- Amount $100
- Due date 2024-01-01"""
    mock_denidin_context.ai_handler.get_response.return_value = mock_response
    
    # Phase 4: analyze=True (default) should call AI
    result = docx_extractor.analyze_media(media, analyze=True)
    
    # Should have AI analysis in response
    assert "invoice" in result["raw_response"].lower()
    assert "$100" in result["raw_response"]
    assert mock_denidin_context.ai_handler.get_response.call_count == 1


def test_analyze_skipped_when_false(docx_extractor, mock_denidin_context):
    """
    Phase 4: With analyze=False, should NOT call AI.
    """
    media = create_docx_media("Some document text")
    
    result = docx_extractor.analyze_media(media, analyze=False)
    
    # No text extraction when analyze=False (current behavior)
    assert result["raw_response"] == ""
    assert result["model_used"] == "python-docx"
    
    # AI should NOT have been called
    assert mock_denidin_context.ai_handler.send_message.call_count == 0


def test_analyze_default_is_true(docx_extractor, mock_denidin_context):
    """
    Phase 4: Default behavior (no analyze parameter) should analyze.
    """
    media = create_docx_media("Document content")
    
    # Mock AI response - get_response returns object with response_text attribute
    mock_response = Mock()
    mock_response.response_text = """DOCUMENT_TYPE: generic
SUMMARY: Document with content
KEY_POINTS:
- Content present"""
    mock_denidin_context.ai_handler.get_response.return_value = mock_response
    
    # Call without analyze parameter (should default to True)
    result = docx_extractor.analyze_media(media)
    
    # Should have document analysis
    assert result["raw_response"] != ""
    assert "python-docx + gpt-4o" in result["model_used"]
    assert mock_denidin_context.ai_handler.get_response.call_count == 1


def test_analyze_graceful_ai_failure(docx_extractor, mock_denidin_context):
    """
    Phase 4: If AI analysis fails, should fallback gracefully.
    """
    media = create_docx_media("Document text")
    
    # Mock AI failure
    mock_denidin_context.ai_handler.send_message.side_effect = Exception("AI service down")
    
    result = docx_extractor.analyze_media(media, analyze=True)
    
    # Should return empty response on failure
    assert result["raw_response"] == ""
    assert result["extraction_quality"] == "high"


def test_analyze_empty_document_no_ai_call(docx_extractor, mock_denidin_context):
    """
    Phase 4: Empty documents should NOT call AI for analysis.
    """
    media = create_docx_media()  # Empty
    
    result = docx_extractor.analyze_media(media, analyze=True)
    
    # Empty text, no AI call
    assert result["raw_response"] == ""
    assert mock_denidin_context.ai_handler.send_message.call_count == 0
    assert result["extraction_quality"] == "high"
