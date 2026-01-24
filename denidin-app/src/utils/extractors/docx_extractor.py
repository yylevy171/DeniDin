"""
DOCX Text Extractor (Feature 003 Phase 3.3 + Phase 4 Enhancement)
Extract text from Word documents using python-docx library.

Phase 4: Optional AI-powered document analysis after text extraction.

CHK Requirements:
- CHK005: Graceful handling of corrupted files
- CHK006: Hebrew text extraction with UTF-8 encoding
- CHK007: Graceful degradation on failures
- CHK010: Layout/structure preservation
- CHK078: Empty document handling
"""
import io
from typing import Dict, List, Optional
from docx import Document
from src.models.media import Media
from src.utils.extractors.base import MediaExtractor


class DOCXExtractor(MediaExtractor):
    """Extract text from DOCX files and optionally analyze with AI."""
    
    def __init__(self, denidin_context):
        """
        Initialize with DeniDin global context.
        
        Args:
            denidin_context: DeniDin instance with ai_handler and config
        """
        super().__init__(denidin_context)
    
    def extract_text(self, media: Media, analyze: bool = True) -> Dict:
        """
        Extract text from DOCX and optionally analyze with AI (Phase 4).
        
        Text extraction is always done via python-docx (no AI needed).
        Document analysis is optional and requires AI call.
        
        CHK006: Hebrew support via UTF-8.
        CHK010: Preserve paragraph structure.
        
        Args:
            media: Media object containing DOCX data in memory
            analyze: If True, use AI to analyze document type/content (Phase 4)
            
        Returns:
            {
                "extracted_text": str,
                "document_analysis": {  # Only if analyze=True
                    "document_type": str,
                    "summary": str,
                    "key_points": List[str]
                } or None,  # None if analyze=False
                "extraction_quality": str,
                "warnings": List[str],
                "model_used": str  # "python-docx" or "python-docx + gpt-4o"
            }
        """
        warnings: List[str] = []
        
        try:
            # Open DOCX from in-memory bytes
            docx_stream = io.BytesIO(media.data)
            doc = Document(docx_stream)
            
            # Extract all paragraph text
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in paragraphs:
                            paragraphs.append(cell_text)
            
            # CHK010: Preserve paragraph structure with double newlines
            extracted_text = "\n\n".join(paragraphs)
            
            # CHK078: Empty document handling
            if not extracted_text:
                warnings.append("Document appears empty")
            
            # Phase 4: Optional AI-powered document analysis
            document_analysis = None
            model_used = "python-docx"
            extraction_quality = "high"  # python-docx is deterministic
            
            if analyze and extracted_text:
                # Call AI to analyze the extracted text
                analysis_result = self._analyze_document(extracted_text)
                document_analysis = analysis_result["document_analysis"]
                model_used = f"python-docx + {analysis_result['model_used']}"
            
            return {
                "extracted_text": extracted_text,
                "document_analysis": document_analysis,
                "extraction_quality": extraction_quality,
                "warnings": warnings,
                "model_used": model_used
            }
            
        except Exception as e:
            # CHK005, CHK007: Graceful failure on corrupted/invalid files
            return {
                "extracted_text": "",
                "document_analysis": None,
                "extraction_quality": "failed",
                "warnings": [f"DOCX extraction failed: {str(e)}"],
                "model_used": "python-docx"
            }
    
    def _analyze_document(self, text: str) -> Dict:
        """
        Analyze extracted text using AI to determine document type and extract insights.
        
        Args:
            text: Extracted text from DOCX
            
        Returns:
            {
                "document_analysis": {
                    "document_type": str,
                    "summary": str,
                    "key_points": List[str]
                },
                "model_used": str
            }
        """
        # Truncate text if too long (to avoid token limits)
        max_chars = 8000
        truncated_text = text[:max_chars]
        if len(text) > max_chars:
            truncated_text += "\n[... text truncated for analysis ...]"
        
        prompt = f"""Analyze this document text and provide:
1. DOCUMENT_TYPE: What type of document is this? (e.g., receipt, invoice, letter, contract, resume, notes, article, generic)
2. SUMMARY: A concise 1-2 sentence summary of the document's content
3. KEY_POINTS: Important information from the document (bullet points)

Document text:
{truncated_text}

Respond in this format:
DOCUMENT_TYPE: <type>
SUMMARY: <summary>
KEY_POINTS:
- <point 1>
- <point 2>
...
"""
        
        try:
            # Use text model for analysis (no vision needed)
            response = self.ai_handler.send_message(
                prompt,
                system_prompt="You are a document analysis assistant. Extract key information clearly and concisely."
            )
            
            # Parse response
            analysis = self._parse_analysis_response(response)
            
            return {
                "document_analysis": analysis,
                "model_used": self.config.ai_model  # Text model, not vision
            }
            
        except Exception as e:
            # Fallback on AI failure
            return {
                "document_analysis": {
                    "document_type": "generic",
                    "summary": "Document analysis failed",
                    "key_points": []
                },
                "model_used": self.config.ai_model
            }
    
    def _parse_analysis_response(self, response: str) -> Dict:
        """
        Parse AI response into structured document_analysis.
        
        Args:
            response: AI response text
            
        Returns:
            {
                "document_type": str,
                "summary": str,
                "key_points": List[str]
            }
        """
        document_type = "generic"
        summary = ""
        key_points = []
        
        lines = response.strip().split("\n")
        current_section = None
        
        for line in lines:
            line = line.strip()
            
            if line.startswith("DOCUMENT_TYPE:"):
                document_type = line.split(":", 1)[1].strip().lower()
            elif line.startswith("SUMMARY:"):
                summary = line.split(":", 1)[1].strip()
            elif line.startswith("KEY_POINTS:"):
                current_section = "key_points"
            elif current_section == "key_points" and line.startswith("-"):
                point = line.lstrip("-").strip()
                if point:
                    key_points.append(point)
        
        return {
            "document_type": document_type,
            "summary": summary or "Document text extracted",
            "key_points": key_points
        }

