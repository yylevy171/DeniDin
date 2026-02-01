"""
DOCX Text Extractor (Feature 003 Phase 3.3 + Phase 4 Enhancement)
Extract text from Word documents using python-docx library.

Phase 4: Optional AI-powered document analysis after text extraction.

CHK Requirements:
- CHK005: Graceful handling of corrupted files
- CHK006: Hebrew text extraction with UTF-8 encoding
- CHK007: Graceful degradation on failures
- CHK010: Layout/structure preservation
- CHK078: Empty document handling
"""
import io
from pathlib import Path
from typing import Dict, List, Optional
import logging
from docx import Document
from src.models.media import Media
from src.handlers.extractors.base import MediaExtractor

logger = logging.getLogger(__name__)


class DOCXExtractor(MediaExtractor):
    """Extract text from DOCX files and optionally analyze with AI."""
    
    def __init__(self, denidin_context):
        """
        Initialize with DeniDin global context.
        
        Args:
            denidin_context: DeniDin instance with ai_handler and config
        """
        super().__init__(denidin_context)
    
    def analyze_media(self, media: Media, analyze: bool = True, caption: str = "") -> Dict:
        """
        Analyze DOCX using AI (Phase 4).
        
        Text extraction is always done via python-docx (no AI needed).
        Document analysis is optional and uses AI call.
        
        CHK006: Hebrew support via UTF-8.
        CHK010: Preserve paragraph structure.
        
        Args:
            media: Media object containing DOCX data in memory
            analyze: If True, use AI to analyze document (Phase 4)
            caption: User's message/question sent with the document (optional)
            
        Returns:
            {
                "raw_response": str,  # AI analysis response
                "extraction_quality": str,
                "warnings": List[str],
                "model_used": str  # "python-docx" or "python-docx + gpt-4o"
            }
        """
        warnings: List[str] = []
        
        try:
            # Open DOCX from in-memory bytes
            docx_stream = io.BytesIO(media.data)
            doc = Document(docx_stream)
            
            # Extract all paragraph text
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in paragraphs:
                            paragraphs.append(cell_text)
            
            # CHK010: Preserve paragraph structure with double newlines
            extracted_text = "\n\n".join(paragraphs)
            
            # CHK078: Empty document handling
            if not extracted_text:
                warnings.append("Document appears empty")
            
            # Phase 4: Optional AI-powered document analysis
            model_used = "python-docx"
            extraction_quality = "high"  # python-docx is deterministic
            raw_response = ""
            
            if analyze and extracted_text:
                # Call AI to analyze the extracted text
                analysis_result = self._analyze_document(extracted_text, caption)
                raw_response = analysis_result.get("raw_response", "")
                model_used = f"python-docx + {analysis_result['model_used']}"
            
            return {
                "raw_response": raw_response,
                "extraction_quality": extraction_quality,
                "warnings": warnings,
                "model_used": model_used
            }
            
        except Exception as e:
            # CHK005, CHK007: Graceful failure on corrupted/invalid files
            return {
                "raw_response": "",
                "extraction_quality": "failed",
                "warnings": [f"DOCX analysis failed: {str(e)}"],
                "model_used": "python-docx"
            }
    
    def _analyze_document(self, text: str, caption: str = "") -> Dict:
        """
        Analyze extracted text using AI to determine document type and extract insights.
        
        Args:
            text: Extracted text from DOCX
            caption: User's message/question sent with the document (optional)
            
        Returns:
            {
                "raw_response": str,  # Full AI response from prompt
                "document_analysis": {
                    "document_type": str,
                    "summary": str,
                    "key_points": List[str]
                },
                "model_used": str
            }
        """
        # Truncate text if too long (to avoid token limits)
        max_chars = 8000
        truncated_text = text[:max_chars]
        if len(text) > max_chars:
            truncated_text += "\n[... text truncated for analysis ...]"
        
        # Build prompt with optional user context
        user_context = f"\n\nUser's question/message: {caption}" if caption else ""
        addressing_note = " addressing the user's question" if caption else ""
        focusing_note = ", focusing on what the user asked about" if caption else ""
        
        # Load prompt template from file (go up 4 levels: extractors → handlers → src → denidin-app)
        prompt_path = Path(__file__).parent.parent.parent.parent / "prompts" / "docx_analysis.txt"
        prompt_template = prompt_path.read_text(encoding="utf-8")
        
        # Format prompt with context
        prompt = prompt_template.format(
            document_text=truncated_text,
            user_context=user_context,
            addressing_note=addressing_note,
            focusing_note=focusing_note
        )
        
        logger.info(f"[DOCXExtractor._analyze_document] Exact prompt being sent ({len(prompt)} chars):")
        logger.info(f"[DOCXExtractor._analyze_document] {prompt}")
        
        try:
            # Load constitution and prepend to prompt (NO system message!)
            constitution = self.ai_handler._load_constitution()
            full_prompt = f"{constitution}\n\n{prompt}" if constitution else prompt
            
            logger.debug(f"[DOCXExtractor._analyze_document] Full prompt length: {len(full_prompt)} chars")
            logger.debug(f"[DOCXExtractor._analyze_document] Constitution loaded: {bool(constitution)}")
            logger.debug(f"[DOCXExtractor._analyze_document] Constitution preview: {constitution[:200] if constitution else 'NONE'}")
            
            # Use text model for analysis (constitution in user prompt, NOT system message)
            from src.models.message import AIRequest
            request = AIRequest(
                user_prompt=full_prompt,
                constitution=constitution,
                max_tokens=self.config.ai_reply_max_tokens,
                temperature=self.config.temperature,
                model=self.config.ai_model,
                chat_id="docx-analysis",
                message_id="docx-analysis"
            )
            ai_response = self.ai_handler.get_response(request)
            response_text = ai_response.response_text

            logger.info(f"[DOCXExtractor._analyze_document] Raw AI response ({len(response_text)} chars):")
            logger.info(f"[DOCXExtractor._analyze_document] {response_text}")

            # No parsing - just pass raw response through as-is
            return {
                "raw_response": response_text,  # Full unmodified AI response
                "document_analysis": {
                    "document_type": "generic",
                    "summary": "See raw_response",
                    "key_points": []
                },
                "model_used": self.config.ai_model  # Text model, not vision
            }
            
        except Exception as e:
            # Fallback on AI failure
            logger.error(f"[DOCXExtractor._analyze_document] Analysis failed: {str(e)}", exc_info=True)
            return {
                "raw_response": "",
                "document_analysis": {
                    "document_type": "generic",
                    "summary": "Document analysis failed",
                    "key_points": []
                },
                "model_used": self.config.ai_model
            }

